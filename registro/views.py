from django.shortcuts import render
from django.utils import timezone
from django.shortcuts import render, get_object_or_404
from django.shortcuts import redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
from .models import Banner_items, solicitud_alumno,Area,Post, Alumno, Proyecto, AsesoresExterno, AsesoresInterno, Empresa, Boleta
from django.http import HttpResponse
import datetime
import os
import glob
import stat
import shutil
import zipfile
from io import StringIO
from io import BytesIO
import pathlib
# Para la generación y descarga del fichero
import io
# Para utilizar algunas de las funciones de la librería
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
import pathlib

#Ruta de la plantilla para los reportes
template_path='registro\\static\\reportes\\the_base.docx'
#Ruta de los archivos temporales
temp_path='registro\\static\\reportes\\temp'
#Ruta del proyecto
PROJECT_PATH = os.path.abspath(os.path.dirname(__name__))

#Pagina principal
def principal(request):
    mensaje=False
    texto=''
    #Si se recarga la página tras rellenar el formulario
    if request.method == "POST":
        #Se obtienen los datos del aaaaaaainforame
        username = request.POST['username']
        matricula = request.POST['matricula']
        carrera = request.POST['carrera']
        mail = request.POST['mail']
        semestre = request.POST['semestre']       
        mensaje=True #Variable de control de notificaciones    
        #Si todos lo campos se han llenado
        if username.strip() and matricula.strip() and carrera.strip() and mail.strip() and semestre.strip():
            #True si ya existe un alumno con la matrícula
            alumno_exists=Alumno.objects.filter(no_control=matricula).exists()
            #True si ya existe una solicitud con la matrícula
            solicitud_exists=solicitud_alumno.objects.filter(s_no_control=matricula).exists()
            if alumno_exists:
                messages.error(request, 'Ya existe un alumno registrado bajo la matrícula enviada.')
            elif solicitud_exists:
                messages.error(request, 'Ya existe una solicitud registrada bajo la matrícula enviada.')
            #Si no existen registros se guarda la solicitud
            else:
                nueva=solicitud_alumno(s_nombre=username,s_no_control=matricula, s_carrera=carrera, s_correo=mail, s_semestre=semestre)
                nueva.save()
                messages.success(request, 'Sus datos han sido enviados correctamente.')
        #Si no se han llenado todos los campos se anexa mensaje de error a la request
        else:
            messages.error(request, 'Debe llenar todos los campos.')
    datos = Empresa.objects.all() #Queryset que retorna todas las empresas (objetos empresa)
    banner_items= Banner_items.objects.order_by('orden_aparicion')
    #Se retorna la página de registro, con o sin mensaje y con los datos de las empresas
    return render(request, 'registro/indexModeloDual.html', {'datos': datos, 'banner_items':banner_items})
#Admin profile
#Value define que sección se mostrará, si solicitudes o proyectos, por defecto es solicitudes
def administrator(request, value=1):
    user = request.user
    #Si el usuario está logueado y es administrador entra, de lo contrario retorna a login
    if request.user.is_authenticated and (user.groups.filter(name='Administradores').exists() or user.is_superuser):
        if value==1:
            return admin_solicitudes(request) #Retorna la vista de solicitudes
        elif value==2:
            return admin_proyectos(request) #Retorna la vista de proyectos_admin
    else:
        return login_general(request)

#Retorna una pagina con todas las solicitudes de alumnos en el perfil de admin
def admin_solicitudes(request):
    #En caso de haber aceptado o rechazado una solicitud
    if request.method == "POST":
        #Si se presionó el boton delete se elimina el registro
        if request.POST.get("delete"):
            matricula = request.POST['no_control']
            solicitud_alumno.objects.filter(s_no_control=matricula).delete()
        #En caso de ser el botón de aceptar el registro se pasan los datos a un nuevo registro en Alumnos y se borra la solicitud
        elif request.POST.get("acept"):
            matricula = request.POST['no_control']
            old=solicitud_alumno.objects.filter(s_no_control=matricula).first() #Obtiene primer resultado de la queryset
            #Se crea un objecto con los datos de la solicitud
            new=Alumno(nombre=old.s_nombre,no_control=old.s_no_control, correo=old.s_correo, semestre=old.s_semestre, carrera=old.s_carrera )
            new.save() #Guarda el Objeto
            old.delete() #Borra el objeto
    #En caso contrario se carga la pagina normalmente con los datos de todas las solicitudes
    datos = solicitud_alumno.objects.all()
    return render(request, 'registro/solicitudes_admin.html', {'datos': datos})

#Muestra todos los proyectos en el perfil de admin
def admin_proyectos(request): 
    proyectos = Proyecto.objects.all() #Queryset que retorna todos los objetos "Proyecto"
    return render(request, 'registro/proyectos_admin.html', {'proyectos': proyectos})

#Carga todos los proyectos y el logo de la empresa a la que está afiliada dicho proyecto
#Retorna la vista de la sección "Proyectos" que aparece en el menu de la pagina principal
def proyectos(request):
    proyectos = Proyecto.objects.all() #Queryset que retorna todos los proyectos
    datos=[] 
    for proyecto in proyectos:
        #Url de la imagen de la empresa del proyecto
        image_url=Empresa.objects.only('url_logo').filter(nombre=proyecto.empresa_rfc).first()
        #Se añade al array un array que contiene el proyecto y la url de la imagen de la empresa
        datos.append([proyecto, image_url])
    #La variable first es un valor de control que se requiere en la plantilla html para mostrar adecuadamente los datos
    return render(request, 'registro/proyectos.html', {'datos': datos,'first': True})

#Retorna la vista con el formulario que permite editar la información del alumno que ha iniciado sesión
def editar_alumno(request):
    user = request.user #Usuario actual en la sesión
    #Entra solo si el usuario está logueado y pertenece al grupo alumnos
    if request.user.is_authenticated and user.groups.filter(name='Alumnos').exists():
        #En caso haber enviado los datos del formulario
        if request.method == "POST":
            #Se obtenen los datos del formulario
            d = request.POST['address']
            t = request.POST['phone']
            c = request.POST['mail']
            r_s = request.POST['red_social']
            r_s_u = request.POST['red_social_url']
            e_a = request.POST['estado_actual']
            #Datos de la contraseña
            old_pass=request.POST['old_pass']
            new_pass=request.POST['new_pass']
            re_new_pass=request.POST['re_new_pass']
            #En caso de no desear cambiar la contraseña solo se actualizan los demás campos
            if not old_pass.strip() and not new_pass.strip() and not re_new_pass.strip():
                Alumno.objects.filter(no_control=request.user.username).update(
                    domicilio=d, telefono=t, correo=c, red_social=r_s, red_social_url=r_s_u,
                    estado_actual=e_a)
                return alumno(request)
            #En caso de querer cambiar la contraseña se valida primero
            elif not old_pass.strip() or not new_pass.strip() or not re_new_pass.strip():
                messages.error(request, 'Si desea cambiar la contraseña debe llenar los tres campos.')
            elif old_pass==new_pass:
                messages.error(request, 'Su contraseña no debe ser igual a la anterior.')
            elif len(new_pass)<6:
                messages.error(request, 'Su contraseña debe tener al menos 6 dígitos.')
            elif new_pass!=re_new_pass:
                messages.error(request, 'Su contraseña y su validación no coinciden.')
            else:
                #Si la validación es exitosa se actualiza el usuario
                Alumno.objects.filter(no_control=request.user.username).update(
                    domicilio=d, telefono=t, correo=c, red_social=r_s, red_social_url=r_s_u,
                    estado_actual=e_a)
                u=User.objects.get(username=request.user.username) #Se obtiene el usuario
                u.set_password(new_pass) #Se define la nueva contraseña
                u.save() #Guardamos cambios
                #Se añade un mensaje de notificación a la request
                messages.error(request, 'Su contraseña ha sido cambiada, por favor inicie sesión de nuevo.')
                #Se cierra sesión y retorna a la pagina principal, solo en caso de haber cambiado la contraseña
                logout(request)
                return render(request, 'registro/indexModeloDual.html')
        #En caso de cargar la pagina normalmente carga los datos que el alumno tenga
        datos = Alumno.objects.filter(no_control=request.user.username).first()
        return render(request, 'registro/editar_alumno.html', {'datos': datos})
    #Si no está logueado retorna a la pagina de login
    else:
        return login_general(request)

#Retorna la vista con el formulario para editar la información de docentes
#Funciona igual que editar_alumno pero adaptado a los campos del modelo AsesoresInterno
def editar_asesor_int(request):
    user = request.user #Usuario actual
    #Entra en caso de existir una sesión inicada de tipo "asesor interno" (Docente)
    if request.user.is_authenticated and user.groups.filter(name='Asesor interno').exists():
        #Se utiliza la misma lógica que con alumnos, pero con campos distintos
        if request.method == "POST":
            #Se obtenen los datos del formulario
            t = request.POST['phone']
            c = request.POST['mail']
            d = request.POST['depto']
            #Datos de la contraseña
            old_pass=request.POST['old_pass']
            new_pass=request.POST['new_pass']
            re_new_pass=request.POST['re_new_pass']
            #En caso de no desear cambiar la contraseña solo se actualizan los demás campos
            if not old_pass.strip() and not new_pass.strip() and not re_new_pass.strip():
                AsesoresInterno.objects.filter(no_empleado=request.user.username).update(
                telefono=t, correo=c, departamento=d)
                return docente(request)
            #En caso de querer cambiar la contraseña se valida primero
            elif not old_pass.strip() or not new_pass.strip() or not re_new_pass.strip():
                messages.error(request, 'Si desea cambiar la contraseña debe llenar los tres campos.')
            elif old_pass==new_pass:
                messages.error(request, 'Su contraseña no debe ser igual a la anterior.')
            elif len(new_pass)<6:
                messages.error(request, 'Su contraseña debe tener al menos 6 dígitos.')
            elif new_pass!=re_new_pass:
                messages.error(request, 'Su contraseña y su validación no coinciden.')
            else:
                #Si la validación es exitosa se actualiza el usuario
                AsesoresInterno.objects.filter(no_empleado=request.user.username).update(
                telefono=t, correo=c, departamento=d)
                u=User.objects.get(username=request.user.username) #Se obtiene el usario django
                u.set_password(new_pass) #Set new pass
                u.save() #Save user updated
                messages.error(request, 'Su contraseña ha sido cambiada, por favor inicie sesión de nuevo.')
                #Se cierra sesión y retorna a la pagina principal, solo en caso de haber cambiado la contraseña
                logout(request)
                return render(request, 'registro/indexModeloDual.html')
        #En caso de cargar la pagina normalmente carga los datos que el alumno tenga
        datos = AsesoresInterno.objects.filter(no_empleado=request.user.username).first()
        departamentos= Area.objects.all()
        return render(request, 'registro/editar_asesor_interno.html', {'datos': datos, 'departamentos': departamentos})
    #Si no está logueado retorna a la pagina de login
    else:
        return login_general(request)

#Retorna la vista con el formulario para editar la información de los asesores externos
#Funciona igual que editar_alumno pero adaptado a los campos del modelo AsesoresExterno
def editar_asesor_ext(request):
    user = request.user #Usuario actual
    #Entra en caso de existir una sesión inicada de tipo "asesor interno" (Docente)
    if request.user.is_authenticated and user.groups.filter(name='Asesor externo').exists():
        #Se utiliza la misma lógica que con alumnos, pero con campos distintos
        if request.method == "POST":
            #Se obtenen los datos del formulario
            t = request.POST['phone']
            c = request.POST['mail']
            #Datos de la contraseña
            old_pass=request.POST['old_pass']
            new_pass=request.POST['new_pass']
            re_new_pass=request.POST['re_new_pass']
            #En caso de no desear cambiar la contraseña solo se actualizan los demás campos
            if not old_pass.strip() and not new_pass.strip() and not re_new_pass.strip():
                AsesoresExterno.objects.filter(no_empleado=request.user.username).update(
                telefono=t, correo=c)
                return asesor_externo(request)
            #En caso de querer cambiar la contraseña se valida primero
            elif not old_pass.strip() or not new_pass.strip() or not re_new_pass.strip():
                messages.error(request, 'Si desea cambiar la contraseña debe llenar los tres campos.')
            elif old_pass==new_pass:
                messages.error(request, 'Su contraseña no debe ser igual a la anterior.')
            elif len(new_pass)<6:
                messages.error(request, 'Su contraseña debe tener al menos 6 dígitos.')
            elif new_pass!=re_new_pass:
                messages.error(request, 'Su contraseña y su validación no coinciden.')
            else:
                #Si la validación es exitosa se actualiza el usuario
                AsesoresExterno.objects.filter(no_empleado=request.user.username).update(
                telefono=t, correo=c)
                u=User.objects.get(username=request.user.username) #Get django user
                u.set_password(new_pass) #Set new pass
                u.save() #Save updated user
                #Append success message, logout user and return to indes page
                messages.error(request, 'Su contraseña ha sido cambiada, por favor inicie sesión de nuevo.')
                logout(request)
                return render(request, 'registro/indexModeloDual.html')
        #En caso de cargar la pagina normalmente carga los datos que el alumno tenga
        datos = AsesoresExterno.objects.filter(no_empleado=request.user.username).first()
        return render(request, 'registro/editar_asesor_externo.html', {'datos': datos})
    #Si no está logueado retorna a la pagina de login
    else:
        return login_general(request)

# Empresas: Retorna la vista Empresas del menu principal
def Empresas(request):
    datos = Empresa.objects.all() #Queryset que retorna los datos de las empresas
    return render(request, 'registro/empresas.html', {'datos': datos, 'first': True})

#Login para todos los usuarios, retorna a su vista correspondiente
def login_general(request):
    #En caso de haber enviado los datos de login
    if request.method == "POST":
        #Se obtienen los datos de la plantilla
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(username=username, password=password) #Se identifica al usuario
        if user is not None:
            if user.is_active:
                logout(request) #En caso de tener otro tipo de sesión abierta esta se cierra
                login(request, user) #Se inicia sesión con los datos ingresados
                if user.groups.filter(name='Alumnos').exists(): #Login alumno
                    return redirect('alumno')
                    #return alumno(request)
                elif user.groups.filter(name='Asesor interno').exists(): #Login docentes
                    return redirect('docente')
                elif user.groups.filter(name='Asesor externo').exists(): #Login asesores externos
                    return redirect('asesor_externo')
                elif user.groups.filter(name='Administradores').exists(): #Login admin
                    return redirect('administrator')
                elif user.is_superuser: #Login admin x2 (admin superuser pero tienen los mismo derechos en su perfil admin)
                    return redirect('administrator')
        #En caso de fallar el login
        else:
            return render(request, "registration/login.html", {'failed': True})
    #En caso de cargar la pagina por primera vez
    return render(request, "registration/login.html")

#Alumno, perfil
def alumno(request):
    user = request.user #Se obtiene al usuario
    #Si esta logueado como alumno entonces
    if request.user.is_authenticated and user.groups.filter(name='Alumnos').exists():
        #Si la pagina se cargó tras guardar el formulario del link de la carpeta de reportes
        if request.method == "POST" and ('url_folder' in request.POST):
            #Obtenemos la nueva url y actualizamos
            report_folder=request.POST['url_folder']
            upload_options=['mega', 'mediafire', 'drive']
            #Si el link enviado no se encuentra en mega, drive o mediafire
            #Se rechazará el envío y se retornará a la carpeta de reportes
            if not any(option in report_folder for option in upload_options):
            	messages.error(request, 'La url debe provenir de uno de los siguientes sitios:'+
            		" Mega, Mediafire, Drive.")
            	return redirect('carpeta_reportes')
            #Caso contrario se aceptará
            Proyecto.objects.filter(alumno=request.user.username).update(
                    url_folder=report_folder)
            messages.success(request, 'Información actualizada.')
        datos = Alumno.objects.filter(no_control=request.user.username).first() #Se obtienen sus datos
        proyecto = Proyecto.objects.filter(alumno=datos.no_control).first() #Los datos de su proyecto si existen
        #Se retorna la pagina con sus datos y los datos de su proyecto
        return render(request, 'registro/alumno.html', {'datos': datos, 'proyecto':proyecto})
    #En caso de no estar logueado se manda a la pagina de login
    else:
        return login_general(request)

# Docentes, perfil
def docente(request):
    user = request.user #Se obtiene al usuario
    #Si está logueado como docente entonces:
    if request.user.is_authenticated and user.groups.filter(name='Asesor interno').exists():
        #En caso de ser redirigido a la pagina tras llena la boleta
        if request.method == "POST" and ('docente_p1' in request.POST):
            #Se obtienen los datos del formulario
            p1 = request.POST['docente_p1']
            p2 = request.POST['docente_p2']
            p3 = request.POST['docente_p3']
            id_boleta = request.POST['id_boleta']
            c = request.POST['comentarios']
            #Se actualizan los campos de la boleta
            boleta_actual=Boleta.objects.get(id=id_boleta)
            boleta_actual.docente_p1=int(p1)
            boleta_actual.docente_p2=int(p2)
            boleta_actual.docente_p3=int(p3)
            boleta_actual.comentarios=c
            boleta_actual.save()
            messages.success(request, "Los datos han sido guardados exitosamente.")
        #Se cargan los datos personales del docente
        datos = AsesoresInterno.objects.filter(no_empleado=request.user.username).first()
        #Se cargan todos los proyectos en los que esté involucrado el docente
        proyectos = Proyecto.objects.filter(asesor_interno=datos.no_empleado)
        #Return docentes.html with perfil and project data
        return render(request, 'registro/docentes.html', {'datos': datos, 'proyectos':proyectos})
    else:
        #En caso de no estar logueado se manda a la pagina de login
        return login_general(request)

#Asesores Externos, perfil
def asesor_externo(request):
    user = request.user #Se obtiene al usuario
    #Si está logueado como docente entonces:
    if request.user.is_authenticated and user.groups.filter(name='Asesor externo').exists():
        #En caso de ser redirigido a la pagina tras llena la boleta
        if request.method == "POST" and ('asesor_p1' in request.POST):
            try:
                #Se obtienen los datos del formulario
                a1 = request.POST['asesor_p1']
                a2 = request.POST['asesor_p2']
                a3 = request.POST['asesor_p3']
                id_boleta = request.POST['id_boleta']
                c = request.POST['comentarios']
                #Se actualizan los campos de la boleta
                boleta_actual=Boleta.objects.get(id=id_boleta)
                boleta_actual.asesor_p1=int(a1)
                boleta_actual.asesor_p2=int(a2)
                boleta_actual.asesort_p3=int(a3)
                boleta_actual.comentarios=c
                boleta_actual.save()
                messages.success(request, "Los datos han sido guardados exitosamente.")
            except:
                #En caso de fallar la operación
                messages.error(request, "Ha ocurrido un error.")
                pass
        #Se cargan los datos personales del asesor externo
        datos = AsesoresExterno.objects.filter(no_empleado=request.user.username).first()
        #Se cargan todos los proyectos en los que esté involucrado el asesor externo
        proyectos = Proyecto.objects.filter(asesor_externo=datos.no_empleado)
        #Return asesoresExt.html with perfil and project data
        return render(request, 'registro/asesoresExt.html', {'datos': datos, 'proyectos':proyectos})
    #En caso de no estar logueado se manda a la pagina de login
    else:
        return login_general(request)

#Boleta con calificaciones de proyecto
def boleta(request):
    user = request.user #Se obtiene al usuario 
    #Si la página fue cargada desde la vista docentes o asesorExt
    #A través del botón "boleta" que es un mini-formulario con envio POST
    #Y llama a esta vista
    if request.method == "POST":
        #Recuperamos la id del proyecto
        id_proyecto = request.POST['id_proyecto']
        if request.user.is_authenticated: #Si el usuario está logueado
            #Como se utiliza una misma plantilla para docentes y asesores
            #Se valida que tipo de usuario y, dependiendo si es:
            #Docente: se solicitan los campos docente_p1, docente_p2, docente_p3 y comentarios
            if user.groups.filter(name='Asesor interno').exists():
                datos=Boleta.objects.filter(proyecto=id_proyecto).only('id','docente_p1', 'docente_p2', 'docente_p3', 'comentarios').first()
                tipo=True
            #Asesir externo: se solicitan los campos asesor_p1, asesor_p2, asesor_p3 y comentarios
            elif user.groups.filter(name='Asesor externo').exists():
                tipo=False
                datos=Boleta.objects.filter(proyecto=id_proyecto).only('id','asesor_p1', 'asesor_p2', 'asesort_p3', 'comentarios').first()
            #Una vez cargados los datos se llama a boletapagina.html y se mandan
            #La variable tipo es un valor de control utilizado en la plantilla
            return render(request, 'registro/boletapagina.html', {'datos': datos, 'tipo':tipo})
        else: #Si no está logueado se retorona
            return login_general(request)
    else:
        return principal(request)

def reportes_administrador(request):
    user = request.user #Se obtiene el usuario actual
    #Si el usuario está logueado y es admin
    if request.user.is_authenticated and (user.groups.filter(name='Administradores').exists() or user.is_superuser):
        #Si el formulario de solicitud de reporte ya se ha llenado
        if request.method == "POST":
            tipo_reporte=request.POST['tipo'] #get data from template
            #Si el checkbox all_dates está seleccionado all_dates=True
            all_dates = True if 'all_dates' in request.POST else False
            if not all_dates:
                #Si el checkbox no está seleccionado se obtiene el rango de fechas
                try:
                    date_first= datetime.datetime.strptime(request.POST['date_one'], '%Y-%m-%d').date()
                    date_last=datetime.datetime.strptime(request.POST['date_two'], '%Y-%m-%d').date()
                except:
                    #En caso de error se asume que serán todas las fechas
                    all_dates=True
            #Se obtienen los datos de todos los proyectos en caso de que all_dates se true
            #En caso de ser false se filtran en base al rango de fechas que especificó el usuario
            proyectos= Proyecto.objects.all() if all_dates else Proyecto.objects.filter(fecha_de_creacion__range=(date_first, date_last)) 
            #Se obtiene la o las empresas seleccionadas
            empresas=request.POST.getlist('empresa')
            #Si se selecciono "todas" o fue alguna de las opciones  all_enterprise se vuelve True
            #Esto nos indica que no se filtrará ninguna empresa
            all_enterprise=True if "todas" in empresas else False
            #En caso de que exista una o varias empresas a filtrar se modifica la consulta previa
            #Dejando solo aquellos proyectos cuya empresa esté dentro de la lista que seleccionó el usuario
            if not all_enterprise:
                proyectos=proyectos.filter(empresa_rfc__in=empresas)

            #El siguiente array guarda arrays del tipo [proyectos_a_guardar_en_un_solo_documento, nombre_documento]
            proyectos_en_secciones=[]
            #Primer tipo de reporte
            if tipo_reporte=="sexo":
                proyectos_h=[]
                proyectos_m=[]
                #Se recorre cada proyecto de la lista filtrada de proyectos y se divide en dos array
                #Cada uno guardando los proyectos donde el estudiante sea hombre o mujer
                for proyecto in proyectos:
                    if proyecto.alumno.sexo=='H' or proyecto.alumno.sexo=='Hombre':
                        proyectos_h.append(proyecto)
                    else:
                        proyectos_m.append(proyecto)
                #Por último se manda a llamar a proyectos en secciones con los datos
                proyectos_en_secciones=[[proyectos_h,'Hombres'], [proyectos_m, 'Mujeres']]            
            #Segundo tipo de reporte
            elif tipo_reporte=="semestre":
                #Por cada i (semestre) se obtienen todos los proyectos de la lista filtrada
                #Donde el alumno de dicho proyecto esté en el semestre i 
                #Se añade esta lista de proyectos filtrados junto su nombre como un elemento de proyectos_en_secciones
                for i in range(1, 13):
                    alumnos=Alumno.objects.filter(semestre=str(i))
                    temp_proyectos=proyectos.filter(alumno__in=alumnos)
                    proyectos_en_secciones.append([temp_proyectos, ('semestre'+str(i))])
            #Tercer tipo de reporte
            elif tipo_reporte=="empresa":
                #De la lista de proyectos filtrada se obtienen las distintas empresas (id)
                empresas=proyectos.values('empresa_rfc').distinct()
                #Por cada empresa en la lista de empresas se obtienen los proyectos
                #realizados en dicha empresa, el nombre de la empresa y se añaden estos
                #valores como un item en proyectos_en_secciones
                for empresa in empresas:
                    temp_proyectos=proyectos.filter(empresa_rfc=empresa['empresa_rfc'])
                    empresa_actual=Empresa.objects.get(rfc_empresa=empresa['empresa_rfc'])
                    proyectos_en_secciones.append([temp_proyectos, empresa_actual.nombre])
            #Cuarto tipo de reporte
            elif tipo_reporte=="asesor_interno":
                #Se obtiene una lista de todos los asesores internos distintos
                #que colaboran en los proyectos filtrados
                docentes=proyectos.values('asesor_interno').distinct()
                #Por cada asesor de la lista se obtienen todos los proyectos de la lista filtrada
                #Donde el docente actual colabore, se añade como un tiem nuevo
                #A proyectos_en_secciones junto con el nombre del docente
                for docente in docentes:
                    temp_proyectos=proyectos.filter(asesor_interno=docente['asesor_interno'])
                    docente_actual=AsesoresInterno.objects.get(no_empleado=docente['asesor_interno'])
                    proyectos_en_secciones.append([temp_proyectos, docente_actual.nombre])
            #Quinto tipo de reporte
            else:
                #Todos los proyectos
                proyectos_en_secciones.append([proyectos, ""])
            #Se llama a reportes y se le envian los proyectos seccionados
            return reportes(request, proyectos_en_secciones)
        empresas= Empresa.objects.all() #Lista de empresas
        #Si es la primera vez que se carga la página se envian las empresas para
        #que estas aparezcan en el select box
        return render(request, 'registro/reportepagina.html', {'empresas':empresas})
    #En caso de no estar logueado o no ser admin retorna a la pagina principal
    #Con su respectivo mensaje de error
    else:
        messages.error(request, 'No ha iniciado sesión como administrador.')
        return principal(request)

def reportes_asesores(request):
    user = request.user #Se obtiene el usuario actual
    #Si el usuario está logueado
    if request.user.is_authenticated:
        #Si el formulario de solicitud de reporte ya se ha llenado
        if request.method == "POST":
            tipo_reporte=request.POST['tipo'] #get data from template
            #Si el checkbox all_dates está seleccionado all_dates=True
            all_dates = True if 'all_dates' in request.POST else False
            if not all_dates:
                try:
                    #Si el checkbox no está seleccionado se obtiene el rango de fechas
                    date_first= datetime.datetime.strptime(request.POST['date_one'], '%Y-%m-%d').date()
                    date_last=datetime.datetime.strptime(request.POST['date_two'], '%Y-%m-%d').date()
                except:
                    #En caso de error se asume que serán todas las fechas
                    all_dates=True
            #Si el usuario que solicita el reporte es un docente
            if user.groups.filter(name='Asesor interno').exists():
                #se obtienen los datos del docente
                a_i=AsesoresInterno.objects.get(no_empleado=request.user.username)
                if all_dates:
                    #Si se quieren todas las fechas se obtienen todos los proyectos
                    #que coincidan con dicho docente
                    proyectos=Proyecto.objects.filter(asesor_interno=a_i)      
                else:
                    #En caso de existir rango de fechas se aplican como restriccion al queryset
                    proyectos=Proyecto.objects.filter(asesor_interno=a_i, fecha_de_creacion__range=(date_first, date_last))
            #En caso de ser un asesor externo quien solicita el reporte
            elif user.groups.filter(name='Asesor externo').exists():
                #Se obtienen sus datos
                a_e=AsesoresExterno.objects.get(no_empleado=request.user.username)
                #Se aplica el mismo filtro de fechas
                if all_dates:
                    proyectos=Proyecto.objects.filter(asesor_externo=a_e)      
                else:
                    proyectos=Proyecto.objects.filter(asesor_externo=a_e, fecha_de_creacion__range=(date_first, date_last))
            #El siguiente array guarda arrays del tipo [proyectos_a_guardar_en_un_solo_documento, nombre_documento]
            proyectos_en_secciones=[]
            #Primer tipo de reporte
            if tipo_reporte=="sexo":
                proyectos_h=[]
                proyectos_m=[]
                #Se recorre cada proyecto de la lista filtrada de proyectos y se divide en dos array
                #Cada uno guardando los proyectos donde el estudiante sea hombre o mujer
                for proyecto in proyectos:
                    if proyecto.alumno.sexo=='H' or proyecto.alumno.sexo=='Hombre':
                        proyectos_h.append(proyecto)
                    else:
                        proyectos_m.append(proyecto)
                #Por último se manda a llamar a proyectos en secciones con los datos
                proyectos_en_secciones=[[proyectos_h,'Hombres'], [proyectos_m, 'Mujeres']]            
            #Segundo tipo de reporte
            elif tipo_reporte=="semestre":
                #Por cada i (semestre) se obtienen todos los proyectos de la lista filtrada
                #Donde el alumno de dicho proyecto esté en el semestre i 
                #Se añade esta lista de proyectos filtrados junto su nombre como un elemento de proyectos_en_secciones
                for i in range(1, 13):
                    alumnos=Alumno.objects.filter(semestre=str(i))
                    temp_proyectos=proyectos.filter(alumno__in=alumnos)
                    proyectos_en_secciones.append([temp_proyectos, ('semestre'+str(i))])
            #Tercer tipo de reporte
            elif tipo_reporte=="empresa":
                #De la lista de proyectos filtrada se obtienen las distintas empresas (id)
                empresas=proyectos.values('empresa_rfc').distinct()
                #Por cada empresa en la lista de empresas se obtienen los proyectos
                #realizados en dicha empresa, el nombre de la empresa y se añaden estos
                #valores como un item en proyectos_en_secciones
                for empresa in empresas:
                    temp_proyectos=proyectos.filter(empresa_rfc=empresa['empresa_rfc'])
                    empresa_actual=Empresa.objects.get(rfc_empresa=empresa['empresa_rfc'])
                    proyectos_en_secciones.append([temp_proyectos, empresa_actual.nombre])
            #Cuarto tipo de reporte
            elif tipo_reporte=='asesor_interno':
                #Se obtiene una lista de todos los asesores internos distintos
                #que colaboran en los proyectos filtrados
                docentes=proyectos.values('asesor_interno').distinct()
                #Por cada asesor de la lista se obtienen todos los proyectos de la lista filtrada
                #Donde el docente actual colabore, se añade como un tiem nuevo
                #A proyectos_en_secciones junto con el nombre del docente
                for docente in docentes:
                    temp_proyectos=proyectos.filter(asesor_interno=docente['asesor_interno'])
                    docente_actual=AsesoresInterno.objects.get(no_empleado=docente['asesor_interno'])
                    proyectos_en_secciones.append([temp_proyectos, docente_actual.nombre])
            #Quinto tipo de reporte: TODOS
            else:
                proyectos_en_secciones.append([proyectos, ""])
            #Se llama a reportes y se le envian los proyectos seccionados
            return reportes(request, proyectos_en_secciones)

        #En caso de ser la primera vez que se ejecuta la página
        #Se cargan la misma plantilla pero con un valor distinto de type
        #Esto permite mostrar diferentes campos para cada tipo de usuario
        if user.groups.filter(name='Asesor interno').exists():
            return render(request, 'registro/reportes_asesores.html', {'type':True})
        elif user.groups.filter(name='Asesor externo').exists():
            return render(request, 'registro/reportes_asesores.html', {'type':False})
    #Si el usuario no está logueado se manda al login
    else:
        return login_general(request)

#Recibe una array de [proyectos, nombre_del_documento]
#Crea los documentos y guarda un array con las rutas de los mismos
def reportes(request, proyectos_seccionados):
    path_files=[] #Ruta de los archivos
    #FOR que recorre los queryset con los proyectos donde cada sección
    #con 1 o más proyectos se vuelve un docx
    for seccion in proyectos_seccionados:
        #Si existe al menos un proyecto en la seccion
        if len(seccion[0])>0:
            #Array que guarda todos los docx creados en la seccion
            compendio_documentos=[]
            #Array que recorre todos los proyectos en la seccion
            #La posicion cero contiene los proyectos, la posicion uno el nombre del documento
            for proyecto in seccion[0]:
                #Por cada proyecto se crea un documento
                document = Document(template_path)
                #Se obtienen el diccionario con los datos a reemplazar  en la plantilla
                data= get_template_data(proyecto.id)
                #Se reemplazan los datos del diccionario en el documento
                docx_replace(document, data)
                #Se guarda el documento en el array
                compendio_documentos.append(document)
            #For que recorre todos los documentos en el compendio de la seccion actual
            #Añade al primer documento todos los demás para terminar con solo 1
            first=True
            for documento in compendio_documentos:
                if first:
                    first=False
                else:
                    #compendio_documentos[0].add_page_break()
                    #Guardamos el contenido del documento actual en el primero de la lista
                    for element in documento.element.body:
                        compendio_documentos[0].element.body.append(element)
            #Se guarda el documento con todos los proyectos de la sección
            #El documento docx se guarda en la ruta temporal
            compendio_documentos[0].save(temp_path+"\\"+"Reporte_"+seccion[1]+".docx")
            #El documento de la sección se añade al array de documentos
            path_files.append("Reporte_"+seccion[1]+".docx")
    #Una vez se han creado y guardado todos los documentos se llama retornar_documentos
    return retornar_documentos(request, path_files)

#Recibe una lista de rutas de archivo, los comprime en un zip y lo descarga desde el navegador
def retornar_documentos(request, path_files):
    #Si no hay ningun documento creado
    if not path_files:
        user = request.user
        messages.error(request, 'No existen registros para descargar.')
        #Retorna una página distinta segun el tipo de sesión iniciada
        if user.groups.filter(name='Asesor interno').exists():
            return render(request, 'registro/reportes_asesores.html', {'type':True})
        elif user.groups.filter(name='Asesor externo').exists():
            return render(request, 'registro/reportes_asesores.html', {'type':False})
        elif user.groups.filter(name='Administradores').exists() or user.is_superuser:
            return redirect('reportes_administrador')
    zipfile_name="Reportes.zip" #Se crea el archivo
    response = HttpResponse(content_type='application/zip')
    zip_file = zipfile.ZipFile(response, 'w')
    #Por cada ruta en el array path_files se guarda un archivo en el zip
    for filename in path_files:
        zip_file.write((PROJECT_PATH+"\\"+temp_path+"\\"+filename), filename)
    #Descar el archivo creado
    response['Content-Disposition'] = 'attachment; filename={}'.format(zipfile_name)
    delete_files_in_folder((PROJECT_PATH+"\\"+temp_path))
    return response

#Descarga el reporte en formato docx del proyecto seleccionado
def reporte_personal(request):
    #Si usó el botón "Generar reporte" que aparece en los proyectos
    if request.method == "POST":
        #Se crea un documento utilizando la plantilla
        document = Document(template_path)
        #Se obtiene la id del proyecto
        id_proyecto=request.POST['id_proyecto']
        #Get project objet
        datos=Proyecto.objects.filter(id=id_proyecto).first()
        #Docx name
        document_name=datos.nombre
        #Obtenemos el dicccionario con los datos del proyecto y sus llaves para reemplazar en la plantilla
        document_content=get_template_data(id_proyecto)
        #Se reemplazan las llaves en la plantilla por los datos del proyecto
        docx_replace(document, document_content)   
        # Save document to memory and download to the user's browser
        document_data = io.BytesIO()
        document.save(document_data)
        document_data.seek(0)
        #Download docs
        response = HttpResponse(
        document_data.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename = "Reporte_"'+document_name+'".docx"'
        response["Content-Encoding"] = "UTF-8"
        delete_files_in_folder((PROJECT_PATH+"\\"+temp_path))
        return response
#Borra los archivos temporales creados, borra el directorio entero y lo vuelve a crear
def delete_files_in_folder(path_folder):
    if os.path.exists(path_folder):
        shutil.rmtree(path_folder)
    os.makedirs(path_folder)
#A partir de una id de proyecto crea y retorna un dccionario con sus valores correspondientes
def get_template_data(id_proyecto):
    datos=Proyecto.objects.get(id=id_proyecto) #Queryset que retorna la info del proyecto
    #Las llaves del diccionario son las variables en el documento docx usado como plantilla
    document_content=dict(
        v_day=datetime.datetime.today().day,
        v_month=datetime.datetime.today().month,
        v_year=datetime.datetime.today().year,
        empresa_name=(datos.empresa_rfc.nombre+"\t"),
        project_name=datos.nombre,
        project_desc=datos.descripcion,
        project_terminated=datos.finalizado,
        project_type=datos.tipo_proyecto,
        project_cal=datos.cal_final,
        student_name=datos.alumno,
        student_no_control=datos.alumno.no_control,
        student_career=datos.alumno.carrera,
        studentsemester=datos.alumno.semestre,
        student_mail=datos.alumno.correo,
        student_phone=datos.alumno.telefono,
        student_address=datos.alumno.domicilio,
        docente_name=datos.asesor_interno,
        docente_depto_name=datos.asesor_interno.departamento.nombre,
        docente_phone=datos.asesor_interno.telefono,
        docente_mail=datos.asesor_interno.correo,
        asesor_name=datos.asesor_externo,
        asesor_phone=datos.asesor_externo.telefono,
        asesor_mail=datos.asesor_externo.correo,
        )
    for key, value in document_content.items():
    	if value is None:
    		document_content[key]="Sin registrar"
    return document_content

#Replace each key of dictionary in docx document with her value
def docx_replace(doc, data):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        for key, val in data.items():
            #key_name = '${{{}}}'.format(key) # I'm using placeholders in the form ${PlaceholderName}
            key_name=key
            if key_name in p.text:
                inline = p.runs
                # Replace strings and retain the same style.
                # The text to be replaced can be split over several runs so
                # search through, identify which runs need to have text replaced
                # then replace the text in those identified
                started = False
                key_index = 0
                # found_runs is a list of (inline index, index of match, length of match)
                found_runs = list()
                found_all = False
                replace_done = False
                for i in range(len(inline)):
                    # case 1: found in single run so short circuit the replace
                    if key_name in inline[i].text and not started:
                        found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                        text = inline[i].text.replace(key_name, str(val))
                        inline[i].text = text
                        replace_done = True
                        found_all = True
                        break
                    if key_name[key_index] not in inline[i].text and not started:
                        # keep looking ...
                        continue

                    # case 2: search for partial text, find first run
                    if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                        # check sequence
                        start_index = inline[i].text.find(key_name[key_index])
                        check_length = len(inline[i].text)
                        for text_index in range(start_index, check_length):
                            if inline[i].text[text_index] != key_name[key_index]:
                            # no match so must be false positive
                                break
                        if key_index == 0:
                            started = True
                        chars_found = check_length - start_index
                        key_index += chars_found
                        found_runs.append((i, start_index, chars_found))
                        if key_index != len(key_name):
                            continue
                        else:
                            # found all chars in key_name
                            found_all = True
                            break
                    # case 2: search for partial text, find subsequent run
                    if key_name[key_index] in inline[i].text and started and not found_all:
                        # check sequence
                        chars_found = 0
                        check_length = len(inline[i].text)
                        for text_index in range(0, check_length):
                            if inline[i].text[text_index] == key_name[key_index]:
                                key_index += 1
                                chars_found += 1
                            else:
                                break
                        # no match so must be end
                        found_runs.append((i, 0, chars_found))
                        if key_index == len(key_name):
                            found_all = True
                            break
                if found_all and not replace_done:
                    for i, item in enumerate(found_runs):
                        index, start, length = [t for t in item]
                        if i == 0:
                            text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                            inline[index].text = text
                        else:
                            text = inline[index].text.replace(inline[index].text[start:start + length], '')
                            inline[index].text = text
                # print(p.text)

#Cierra sesión y retorna a la página principal
def logout_user(request):
    logout(request)
    return redirect('principal')

#Retorna la vista que permite ingresar el link donde están subidos los reportes
def carpeta_reportes(request):
    user = request.user #Se obtiene al usuario
    #Si esta logueado como alumno entonces
    if request.user.is_authenticated and user.groups.filter(name='Alumnos').exists():
        #Se envian los datos del proyecto del alumno
        try:
            datos=Proyecto.objects.get(alumno=request.user.username)
        except:
            datos=Proyecto.objects.filter(alumno=request.user.username).first()
        return render(request, 'registro/carpeta_reportes.html', {'datos':datos})
    else:
        #Si no es un alumno se va al login general
        return login_general(request)
def album_fotos(request):
    images=Post.objects.all()
    return render(request, 'registro/album_fotos.html', {'images':images})


